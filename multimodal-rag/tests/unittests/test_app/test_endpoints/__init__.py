from fastapi.testclient import TestClient
from mm_rag.api.main import app
from PIL import Image
from reportlab.pdfgen import canvas
from docx import Document
from fastapi import Response
from fastapi.responses import JSONResponse

from io import BufferedReader
from mm_rag.utils import get_secret


test_client = TestClient(app)


def create_test_pdf(filename="test.pdf"):
  # Create a canvas object with the given filename
  c = canvas.Canvas(filename)

  # Add some minimal text
  c.drawString(100, 750, "Hello, this is a test PDF.")

  # Finalize the PDF file
  c.save()
  print(f"PDF created: {filename}")


def create_docx(path: str):
  doc = Document()

  doc.add_heading("TestDoc", 0)
  doc.add_paragraph(
    "Test Paragraph"
  )
  doc.add_page_break()
  table = doc.add_table(2, 2)
  table.cell(0, 0).text = "Test"
  table.cell(0, 1).text = "Row"
  table.cell(1, 0).text = "Another"
  table.cell(1, 1).text = "Cell"

  doc.save(path)


def send_file_request(file_path: str, content: BufferedReader | bytes, token: str,
                 content_type: str = "multipart/form-data") -> Response | JSONResponse:
  def _send(name: str, file_content: bytes) -> Response:
    return test_client.post(
      '/upload-file',
      headers={"Authorization": f"Bearer {token}"},
      files={"file": (name, file_content, content_type)}
    )

  return _send(
    name=file_path,
    file_content=content
  )


def write_to_file(file_path: str, content: str | bytes) -> None:
  if isinstance(content, str):
    with open(file_path, 'w') as f:
      f.write(content)
  elif isinstance(content, bytes):
    with open(file_path, 'wb') as f:
      f.write(content)


def create_img(path: str, size: tuple[int, int] = (100, 100)) -> None:
  img = Image.new('RGB', size)
  img.save(path, 'JPEG')


def send_chat_request(prompt: str) -> Response | JSONResponse:
  def _send(p: str) -> Response:
    return test_client.post(
      '/chat',
      headers={"Authorization": f"Bearer {get_secret()['bearer_pat']}"},
      json={"query": p}
    )

  return _send(prompt)


def send_search_request(query: str) -> Response | JSONResponse:
  return test_client.post(
      '/search',
      headers={"Authorization": f"Bearer {get_secret()['bearer_pat']}"},
      json={"query": query}
    )

def send_cleanup_request(token: str) -> Response | JSONResponse:
  return test_client.post(
    '/cleanUp',
    headers={"Authorization": f"Bearer {token}"},
  )
