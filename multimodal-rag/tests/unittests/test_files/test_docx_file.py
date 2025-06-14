import unittest
from unittest.mock import MagicMock, patch

from mm_rag.processing.files import DocxFile
from mm_rag.processing.processors import DocxProcessor
from mm_rag.processing.handlers import ImgHandler

import os
from docx import Document

from PIL import Image


FILE_PATH = "test.docx"


def create_doc(file_path: str):
  doc = Document()

  doc.add_heading("TestDoc", 0)
  doc.add_paragraph(
      "Test Paragraph"
  )
  doc.add_page_break()
  table = doc.add_table(2, 2)
  table.cell(0, 0).text = "Test"
  table.cell(0, 1).text = "Row"
  table.cell(1, 0).text = "Another"
  table.cell(1, 1).text = "Cell"

  doc.save(file_path)


class TestDocxFile(unittest.TestCase):
  def setUp(self):
    self.mock_processor = DocxProcessor(MagicMock(), MagicMock(ImgHandler))
    create_doc(FILE_PATH)
    self.file = DocxFile(
      FILE_PATH, 'user123', self.mock_processor
    )

  def tearDown(self):
    new_path, _ = os.path.splitext(FILE_PATH)
    new_path = new_path + '.pdf'

    if os.path.exists(FILE_PATH):
      os.remove(FILE_PATH)

    if os.path.exists(new_path):
      os.remove(new_path)

  def test_right_file_creation(self):
    file_content = self.file.file_content

    self.assertIsInstance(
      file_content, list
    )
    for page in file_content:
      self.assertIsInstance(
        page, Image.Image
      )

    self.assertTrue(
      self.file.file_path.endswith('.docx')
    )

  def test_right_encodings(self):
    expected = 'encoded'
    self.mock_processor.handler.base64_encode.return_value = str(expected)

    self.assertIsInstance(self.file.encodings, list)
    for encoded in self.file.encodings:
      self.assertIsInstance(encoded, str)
      self.assertEqual(
        encoded, expected
      )

  def test_right_ids(self):
    self.assertEqual(
      self.file.file_id, self.file.metadata.fileId
    )


if __name__ == "__main__":
  unittest.main()